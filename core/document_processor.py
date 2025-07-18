"""
Document processing core logic.

This module contains the main business logic for document processing,
text extraction, and embedding generation.
"""
import logging
import tempfile
import os
import hashlib
from typing import Dict, Any, Optional, List
from pathlib import Path
from datetime import datetime

from services.azure_blob_service import azure_blob_service
from services.openai_service import openai_service
from services.redis_service import redis_service
from services.computer_vision_service import computer_vision_service

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Core document processing logic."""
    
    def __init__(
        self,
        blob_service=None,
        openai_service=None,
        redis_service=None,
        vision_service=None,
        logger_instance=None
    ):
        """Initialize document processor with optional dependency injection."""
        from services.azure_blob_service import azure_blob_service as _blob
        from services.openai_service import openai_service as _openai
        from services.redis_service import redis_service as _redis
        from services.computer_vision_service import computer_vision_service as _vision
        self.blob_service = blob_service or _blob
        self.openai_service = openai_service or _openai
        self.redis_service = redis_service or _redis
        self.vision_service = vision_service or _vision
        self.logger = logger_instance or logger
    
    def process_document_from_queue(
        self, 
        blob_name: str, 
        blob_url: str = "", 
        file_size: int = 0, 
        content_type: str = ""
    ) -> bool:
        """
        Process document from queue message.
        
        Args:
            blob_name: Name of the blob to process
            blob_url: URL of the blob (optional)
            file_size: Size of the file (optional)
            content_type: MIME content type (optional)
            
        Returns:
            True if processing was successful, False otherwise
        """
        try:
            self.logger.info(f"[TRACE] Iniciando procesamiento de documento desde cola: {blob_name}")
            # Create temporary file for processing
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(blob_name).suffix) as temp_file:
                temp_file_path = temp_file.name
            
            try:
                # Download file from blob storage
                download_success = self.blob_service.download_file(blob_name, temp_file_path)
                if not download_success:
                    self.logger.error(f"Failed to download blob: {blob_name}")
                    return False
                
                self.logger.info(f"Successfully downloaded blob: {blob_name}")
                
                # Get file metadata
                file_metadata = self.blob_service.get_blob_metadata(blob_name) or {}
                
                # Calculate file hash for document ID
                file_hash = self._calculate_file_hash(temp_file_path)
                document_id = self._generate_document_id(blob_name, file_hash)
                
                # Extract text based on file type
                extracted_text = self._extract_text_from_file(temp_file_path, blob_name, content_type)
                if not extracted_text:
                    self.logger.warning(f"No text extracted from file: {blob_name}")
                    return False
                
                self.logger.info(f"[TRACE] Extrayendo texto del archivo: {blob_name}")
                
                # Clean and chunk text
                cleaned_text = self._clean_text(extracted_text)
                text_chunks = self._chunk_text(cleaned_text, chunk_size=1000, overlap=100)
                
                self.logger.info(f"Text extracted and chunked. Original length: {len(extracted_text)}, Chunks: {len(text_chunks)}")
                self.logger.info(f"[TRACE] Limpiando y dividiendo el texto extraído en chunks")
                
                # Generate embeddings for each chunk
                embeddings = self._generate_embeddings_for_chunks(text_chunks)
                if not embeddings:
                    self.logger.error("No embeddings generated for any chunks")
                    return False
                
                self.logger.info(f"[TRACE] Generando embeddings para los chunks del documento")
                
                # Store embeddings and metadata in Redis
                self._store_document_embeddings(document_id, blob_name, embeddings, file_metadata)
                
                self.logger.info(f"[TRACE] Guardando embeddings y metadatos en Redis para el documento: {document_id}")
                
                # Update blob metadata to mark as processed
                self._update_blob_metadata(blob_name, document_id, len(embeddings))
                
                self.logger.info(f"[TRACE] Actualizando metadatos del blob para el documento procesado: {document_id}")
                
                self.logger.info(f"Successfully processed document: {document_id}")
                return True
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    self.logger.info(f"Temporary file cleaned up: {temp_file_path}")
                    
        except Exception as e:
            self.logger.error(f"Failed to process document {blob_name}: {e}")
            return False
    
    def process_document_from_blob(self, blob_stream, blob_name: str) -> bool:
        """
        Procesa archivos subidos a Blob Storage: texto, imágenes y PDF (texto o escaneados).
        Divide en chunks, genera embeddings por chunk y guarda cada uno en Redis.
        """
        self.logger.info(f"[DEBUG] Iniciando process_document_from_blob para: {blob_name}")
        try:
            # Crear archivo temporal
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(blob_name).suffix) as temp_file:
                temp_file.write(blob_stream.read())
                temp_file_path = temp_file.name
                self.logger.info(f"[DEBUG] Archivo temporal creado: {temp_file_path}")
            try:
                # Obtener metadata
                file_metadata = self.blob_service.get_blob_metadata(blob_name) or {}
                self.logger.info(f"[DEBUG] Metadata del archivo obtenida: {file_metadata}")
                # Detectar tipo de archivo
                ext = Path(blob_name).suffix.lower()
                extracted_text = ""
                if ext in ['.txt', '.md', '.csv']:
                    self.logger.info(f"[DEBUG] Procesando como archivo de texto: {blob_name}")
                    with open(temp_file_path, 'r', encoding='utf-8') as f:
                        extracted_text = f.read()
                elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff']:
                    self.logger.info(f"[DEBUG] Procesando como imagen: {blob_name}")
                    with open(temp_file_path, 'rb') as f:
                        image_bytes = f.read()
                    extracted_text = self.vision_service.extract_text_from_image_bytes(image_bytes)
                elif ext == '.pdf':
                    self.logger.info(f"[DEBUG] Procesando como PDF: {blob_name}")
                    # Intentar extraer texto directo
                    extracted_text = self._extract_text_from_file(temp_file_path, blob_name)
                    if not extracted_text or len(extracted_text.strip()) < 20:
                        self.logger.info(f"[DEBUG] PDF parece escaneado, usando visión computacional")
                        # Extraer imágenes de cada página y pasar por visión
                        try:
                            import fitz  # PyMuPDF
                            doc = fitz.open(temp_file_path)
                            pdf_text = []
                            for page_num in range(len(doc)):
                                page = doc.load_page(page_num)
                                images = page.get_images(full=True)
                                for img_index, img in enumerate(images):
                                    xref = img[0]
                                    base_image = doc.extract_image(xref)
                                    image_bytes = base_image['image']
                                    text_img = self.vision_service.extract_text_from_image_bytes(image_bytes)
                                    if text_img:
                                        pdf_text.append(text_img)
                            extracted_text = '\n'.join(pdf_text)
                        except Exception as e:
                            self.logger.error(f"[ERROR] Error extrayendo texto de PDF escaneado: {e}")
                else:
                    self.logger.warning(f"[WARNING] Formato de archivo no soportado: {blob_name}")
                    return False
                if not extracted_text or len(extracted_text.strip()) == 0:
                    self.logger.warning(f"[WARNING] No se extrajo texto del archivo: {blob_name}")
                    return False
                self.logger.info(f"[DEBUG] Texto extraído exitosamente. Longitud: {len(extracted_text)} caracteres")
                # Chunking
                self.logger.info(f"[DEBUG] Iniciando chunking del texto")
                text_chunks = self._chunk_text(extracted_text, chunk_size=1000, overlap=100)
                self.logger.info(f"[DEBUG] Texto dividido en {len(text_chunks)} chunks")
                # Procesar cada chunk: generar embedding y guardar en Redis
                for idx, chunk in enumerate(text_chunks):
                    self.logger.info(f"[DEBUG] Procesando chunk {idx+1}/{len(text_chunks)}")
                    embedding = self.openai_service.generate_embedding(chunk)
                    if embedding:
                        chunk_metadata = {
                            "chunk_index": idx,
                            "document_id": blob_name,
                            "filename": blob_name,
                            "content": chunk,
                            "embedding": embedding,
                            "file_metadata": file_metadata
                        }
                        self.redis_service.set_cache(f"doc_chunk:{blob_name}:{idx}", chunk_metadata)
                        self.logger.info(f"[SUCCESS] Chunk {idx+1} almacenado en Redis para {blob_name}")
                    else:
                        self.logger.error(f"[ERROR] No se pudo generar embedding para chunk {idx+1} de {blob_name}")
                # Actualizar metadatos del blob
                self.logger.info(f"[DEBUG] Actualizando metadatos del blob")
                self._update_blob_metadata(blob_name, blob_name, len(text_chunks))
                self.logger.info(f"[DEBUG] Metadatos del blob actualizados")
                self.logger.info(f"[SUCCESS] Documento procesado exitosamente: {blob_name}")
                return True
            finally:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    self.logger.info(f"[DEBUG] Archivo temporal eliminado: {temp_file_path}")
        except Exception as e:
            self.logger.error(f"[ERROR] Error procesando blob {blob_name}: {e}")
            import traceback
            self.logger.error(f"[ERROR] Traceback completo: {traceback.format_exc()}")
            return False
    
    def process_document(self, file_url: str, file_name: str, user_id: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """
        Process document from URL with options.
        
        Args:
            file_url: URL of the document to process
            file_name: Name of the file
            user_id: User ID for tracking
            options: Processing options (extract_text, generate_embeddings, etc.)
            
        Returns:
            Dictionary with processing results
        """
        try:
            # Validate file URL
            if not self.validate_file_url(file_url):
                return {
                    "success": False,
                    "error": "Invalid file URL",
                    "processing_id": f"proc_{user_id}_{int(datetime.now().timestamp())}"
                }
            
            # Get file extension and validate format
            file_extension = self.get_file_extension(file_name)
            if not self.is_supported_format(file_extension):
                return {
                    "success": False,
                    "error": f"Unsupported file format: {file_extension}",
                    "processing_id": f"proc_{user_id}_{int(datetime.now().timestamp())}"
                }
            
            # Download file content
            import requests
            response = requests.get(file_url, timeout=30)
            response.raise_for_status()
            
            # Extract text based on file type
            extracted_text = ""
            if options.get("extract_text", True):
                if file_extension in ['jpg', 'jpeg', 'png', 'bmp', 'gif', 'tiff']:
                    extracted_text = self.vision_service.extract_text_from_image_bytes(response.content)
                elif file_extension == 'pdf':
                    extracted_text = self._extract_text_from_pdf_bytes(response.content)
                elif file_extension in ['docx', 'doc']:
                    extracted_text = self._extract_text_from_word_bytes(response.content)
                elif file_extension in ['txt', 'md', 'csv']:
                    extracted_text = response.text
                else:
                    return {
                        "success": False,
                        "error": f"Text extraction not supported for format: {file_extension}",
                        "processing_id": f"proc_{user_id}_{int(datetime.now().timestamp())}"
                    }
            
            # Generate embeddings if requested
            embeddings = None
            if options.get("generate_embeddings", False) and extracted_text:
                embeddings = self.openai_service.generate_embedding(extracted_text)
            
            # Build result
            result = {
                "success": True,
                "content": extracted_text,
                "embeddings": embeddings,
                "metadata": {
                    "file_name": file_name,
                    "file_type": file_extension,
                    "user_id": user_id,
                    "processing_time": datetime.now().isoformat()
                },
                "processing_id": f"proc_{user_id}_{int(datetime.now().timestamp())}"
            }
            
            # Add file size warning if large
            if len(response.content) > 10 * 1024 * 1024:  # 10MB
                result["metadata"]["warning"] = "Large file detected"
            
            return result
            
        except requests.RequestException as e:
            return {
                "success": False,
                "error": f"Failed to download file: {str(e)}",
                "processing_id": f"proc_{user_id}_{int(datetime.now().timestamp())}"
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"Failed to process document: {str(e)}",
                "processing_id": f"proc_{user_id}_{int(datetime.now().timestamp())}"
            }

    def validate_file_url(self, url: str) -> bool:
        """
        Validate if URL is a valid file URL.
        
        Args:
            url: URL to validate
            
        Returns:
            True if valid file URL
        """
        try:
            if not url or not isinstance(url, str):
                return False
            
            # Basic URL validation
            if not url.startswith(('http://', 'https://')):
                return False
            
            # Check for invalid characters
            if ' ' in url or '\n' in url or '\t' in url:
                return False
                
            return True
        except Exception:
            return False

    def get_file_extension(self, filename: str) -> str:
        """
        Extract file extension from filename.
        
        Args:
            filename: Name of the file
            
        Returns:
            File extension (lowercase, without dot)
        """
        try:
            if not filename:
                return ""
            
            # Split by dots and get the last part
            parts = filename.split('.')
            if len(parts) <= 1:
                return ""
            
            return parts[-1].lower()
        except Exception:
            return ""

    def is_supported_format(self, file_extension: str) -> bool:
        """
        Check if file format is supported.
        
        Args:
            file_extension: File extension to check
            
        Returns:
            True if format is supported
        """
        supported_formats = ['pdf', 'docx', 'doc', 'txt', 'md', 'csv', 'jpg', 'jpeg', 'png', 'bmp', 'gif', 'tiff']
        return file_extension.lower() in supported_formats

    def _extract_text_from_file(self, file_path: str, blob_name: str, content_type: str = "") -> str:
        """
        Extract text from file based on its type and content type.
        
        Args:
            file_path: Path to the temporary file
            blob_name: Original blob name for reference
            content_type: MIME content type of the file
            
        Returns:
            Extracted text content
        """
        try:
            file_extension = Path(blob_name).suffix.lower()
            
            # Check if it's an image file (based on extension or content type)
            if (file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff'] or 
                'image/' in content_type.lower()):
                self.logger.info(f"Processing image file with OCR: {blob_name}")
                return self.vision_service.extract_text_from_image_file(file_path)
                
            elif file_extension == '.pdf' or content_type == 'application/pdf':
                self.logger.info(f"Processing PDF file: {blob_name}")
                return self._extract_text_from_pdf(file_path)
                
            elif file_extension in ['.docx', '.doc'] or 'word' in content_type.lower():
                self.logger.info(f"Processing Word document: {blob_name}")
                return self._extract_text_from_word(file_path)
                
            elif (file_extension in ['.txt', '.md', '.csv'] or 
                  'text/' in content_type.lower() or 
                  'plain' in content_type.lower()):
                self.logger.info(f"Processing text file: {blob_name}")
                return self._extract_text_from_text_file(file_path)
                
            else:
                self.logger.warning(f"Unsupported file type: {file_extension} (content_type: {content_type})")
                return ""
                
        except Exception as e:
            self.logger.error(f"Failed to extract text from {blob_name}: {e}")
            raise
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            from pypdf import PdfReader
            
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            return text.strip()
            
        except Exception as e:
            self.logger.error(f"Failed to extract text from PDF: {e}")
            raise
    
    def _extract_text_from_word(self, file_path: str) -> str:
        """Extract text from Word document."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
            
        except Exception as e:
            self.logger.error(f"Failed to extract text from Word document: {e}")
            raise
    
    def _extract_text_from_text_file(self, file_path: str) -> str:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
                
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                self.logger.error(f"Failed to read text file with alternative encoding: {e}")
                raise
        except Exception as e:
            self.logger.error(f"Failed to read text file: {e}")
            raise
    
    def _generate_embeddings_for_chunks(self, text_chunks: List[str]) -> List[Dict[str, Any]]:
        """
        Generate embeddings for text chunks.
        
        Args:
            text_chunks: List of text chunks to process
            
        Returns:
            List of embeddings with metadata
        """
        embeddings = []
        for i, chunk in enumerate(text_chunks):
            try:
                embedding = self.openai_service.generate_embedding(chunk)
                if embedding:
                    embeddings.append({
                        "chunk_index": i,
                        "text": chunk,
                        "embedding": embedding
                    })
                    self.logger.info(f"Generated embedding for chunk {i+1}/{len(text_chunks)}")
                else:
                    self.logger.error(f"Failed to generate embedding for chunk {i}")
            except Exception as e:
                self.logger.error(f"Failed to generate embedding for chunk {i}: {e}")
                continue
        
        return embeddings
    
    def _store_document_embeddings(
        self, 
        document_id: str, 
        blob_name: str, 
        embeddings: List[Dict[str, Any]], 
        file_metadata: Dict[str, Any]
    ) -> None:
        """
        Store document embeddings and metadata in Redis.
        
        Args:
            document_id: Unique document identifier
            blob_name: Original blob name
            embeddings: List of embeddings with text chunks
            file_metadata: File metadata from blob storage
        """
        self.logger.info(f"[DEBUG] Iniciando _store_document_embeddings para documento: {document_id}")
        try:
            # Create document metadata
            document_metadata = {
                "document_id": document_id,
                "filename": blob_name,
                "text": " ".join([emb["text"] for emb in embeddings]),
                "content_type": file_metadata.get("content_type", ""),
                "upload_date": file_metadata.get("upload_date", ""),
                "file_size": file_metadata.get("file_size", 0),
                "chunks_count": len(embeddings),
                "processing_timestamp": file_metadata.get("processing_timestamp", ""),
                "embeddings_generated": "true"
            }
            self.logger.info(f"[DEBUG] Metadata del documento creada: {document_metadata}")
            
            # Store main document embedding (average of all chunks)
            if embeddings:
                # Calculate average embedding
                avg_embedding = []
                embedding_length = len(embeddings[0]["embedding"])
                self.logger.info(f"[DEBUG] Calculando embedding promedio de {len(embeddings)} embeddings de longitud {embedding_length}")
                
                for i in range(embedding_length):
                    avg_value = sum(emb["embedding"][i] for emb in embeddings) / len(embeddings)
                    avg_embedding.append(avg_value)
                
                self.logger.info(f"[DEBUG] Embedding promedio calculado. Longitud: {len(avg_embedding)}")
                
                # Store in Redis
                self.logger.info(f"[DEBUG] Intentando almacenar en Redis con clave: embedding:{document_id}")
                redis_success = self.redis_service.store_embedding(document_id, avg_embedding, document_metadata)
                if redis_success:
                    self.logger.info(f"[SUCCESS] Embedding almacenado exitosamente en Redis para documento: {document_id}")
                else:
                    self.logger.error(f"[ERROR] Falló el almacenamiento en Redis para documento: {document_id}")
                
        except Exception as e:
            self.logger.error(f"[ERROR] Error en _store_document_embeddings para {document_id}: {e}")
            import traceback
            self.logger.error(f"[ERROR] Traceback completo: {traceback.format_exc()}")
            raise
    
    def _update_blob_metadata(self, blob_name: str, document_id: str, chunks_count: int) -> None:
        """
        Update blob metadata to mark embeddings as added.
        
        Args:
            blob_name: Name of the blob
            document_id: Generated document ID
            chunks_count: Number of text chunks processed
        """
        try:
            metadata = {
                "processed": "true",
                "document_id": document_id,
                "chunks_count": str(chunks_count),
                "embeddings_generated": "true",
                "processed_timestamp": datetime.now().isoformat()
            }
            
            # Update blob metadata
            self.blob_service.update_blob_metadata(blob_name, metadata)
            self.logger.info(f"Updated blob metadata for {blob_name}")
            
        except Exception as e:
            self.logger.error(f"Failed to update blob metadata: {e}")
            raise
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA-256 hash of file."""
        try:
            hash_sha256 = hashlib.sha256()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_sha256.update(chunk)
            return hash_sha256.hexdigest()
        except Exception as e:
            self.logger.error(f"Failed to calculate file hash: {e}")
            raise
    
    def _generate_document_id(self, blob_name: str, file_hash: str) -> str:
        """Generate unique document ID."""
        try:
            # Create a unique document ID based on filename and hash
            base_name = Path(blob_name).stem
            return f"{base_name}_{file_hash[:8]}"
        except Exception as e:
            self.logger.error(f"Failed to generate document ID: {e}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        try:
            # Remove extra whitespace
            text = ' '.join(text.split())
            # Remove special characters that might cause issues
            text = text.replace('\x00', '')
            return text.strip()
        except Exception as e:
            self.logger.error(f"Failed to clean text: {e}")
            return text
    
    def _chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 100) -> List[str]:
        """Split text into overlapping chunks."""
        try:
            if len(text) <= chunk_size:
                return [text]
            
            chunks = []
            start = 0
            
            while start < len(text):
                end = start + chunk_size
                chunk = text[start:end]
                chunks.append(chunk)
                start = end - overlap
                
                if start >= len(text):
                    break
            
            return chunks
        except Exception as e:
            self.logger.error(f"Failed to chunk text: {e}")
            return [text] 

    def _extract_text_from_pdf_bytes(self, pdf_bytes: bytes) -> str:
        """Extract text from PDF bytes."""
        try:
            from pypdf import PdfReader
            from io import BytesIO
            
            text = ""
            pdf_stream = BytesIO(pdf_bytes)
            pdf_reader = PdfReader(pdf_stream)
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            self.logger.error(f"Failed to extract text from PDF bytes: {e}")
            raise

    def _extract_text_from_word_bytes(self, docx_bytes: bytes) -> str:
        """Extract text from Word document bytes."""
        try:
            from docx import Document
            from io import BytesIO
            
            doc_stream = BytesIO(docx_bytes)
            doc = Document(doc_stream)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
        except Exception as e:
            self.logger.error(f"Failed to extract text from Word document bytes: {e}")
            raise 